#!/usr/bin/env python3
"""Strategic Counsel v3.3 - CH Year Order, Export/Memory, Protocol/Red Flag UI

Key Changes:
- CH Summaries now automatically use a Gemini model (via ch_pipeline.py logic).
- Sidebar AI model selector is now only for 'Consult Counsel & Digest Updates'.
- CH Pipeline returns AI summarization costs, displayed in UI.
- Implemented more accurate token counting for Gemini in 'Consult Counsel'.
- Corrected attribute access for Gemini SDK check.
- CH Results display uses st.expander per company.
- Added UI for keyword-based filtering in CH analysis (backend logic placeholder).
- Added "Copy Summary" to CH expanders (via st.code).
- CH Summaries can now be selected for injection into Counsel chat.
- Added Protocol status display in sidebar.
- Attempt to highlight "Red Flags" section from CH summaries.
"""

from __future__ import annotations

import streamlit as st
st.set_page_config(
    page_title="Strategic Counsel", page_icon="⚖️", layout="wide",
    initial_sidebar_state="expanded",
    menu_items={'About': "# Strategic Counsel v3.3\nModular AI Legal Assistant Workspace."}
)

try:
    import config
    logger = config.logger
    from ch_pipeline import TEXTRACT_AVAILABLE as CH_PIPELINE_TEXTRACT_FLAG_FROM_MODULE
    config.CH_PIPELINE_TEXTRACT_FLAG = CH_PIPELINE_TEXTRACT_FLAG_FROM_MODULE
except ImportError as e_initial_imports:
    st.error(f"Fatal Error: Could not import core modules (config, ch_pipeline): {e_initial_imports}")
    st.stop()
except Exception as e_conf:
    st.error(f"Fatal Error during config.py import or setup: {e_conf}")
    st.stop()

import datetime as _dt
import hashlib as _hashlib
import io
import json
import os
import pathlib as _pl
import re # For Red Flag parsing
import tempfile
import csv
from typing import List, Tuple, Dict, Optional, Any

import pandas as pd
from docx import Document

try:
    from app_utils import (
        summarise_with_title, fetch_url_content, find_company_number,
        extract_text_from_uploaded_file
    )
    from about_page import render_about_page
    from ch_pipeline import run_batch_company_analysis
    from ai_utils import get_improved_prompt # Added import
    # from ai_utils import _gemini_generate_content_with_retry_and_tokens # Not directly used in app.py typically
except ImportError as e_app_utils_more:
    st.error(f"Fatal Error: Could not import app utilities or CH pipeline: {e_app_utils_more}")
    logger.error(f"ImportError from app_utils/about_page/ch_pipeline/ai_utils: {e_app_utils_more}", exc_info=True)
    st.stop()

APP_BASE_PATH: _pl.Path = config.APP_BASE_PATH
OPENAI_API_KEY_PRESENT = bool(config.OPENAI_API_KEY and config.OPENAI_API_KEY.startswith("sk-"))
CH_API_KEY_PRESENT = bool(config.CH_API_KEY)
GEMINI_API_KEY_PRESENT = bool(config.GEMINI_API_KEY and config.genai) # Corrected to check config.genai

REQUIRED_DIRS_REL = ("memory", "memory/digests", "summaries", "exports", "logs", "static")
for rel_p in REQUIRED_DIRS_REL:
    abs_p = APP_BASE_PATH / rel_p
    try: abs_p.mkdir(parents=True, exist_ok=True)
    except OSError as e_mkdir: st.error(f"Fatal Error creating directory {abs_p.name}: {e_mkdir}"); st.stop()

MODEL_PRICES_PER_1K_TOKENS_GBP: Dict[str, float] = {
    "gpt-4o": 0.0040, "gpt-4-turbo": 0.0080, "gpt-3.5-turbo": 0.0004, "gpt-4o-mini": 0.00012,
    config.GEMINI_MODEL_DEFAULT: 0.0028,
    "gemini-1.5-pro-latest": 0.0028, # Ensure this matches config.GEMINI_MODEL_DEFAULT if it's the one
    "gemini-1.5-flash-latest": 0.00028
}
MODEL_ENERGY_WH_PER_1K_TOKENS: Dict[str, float] = {
    "gpt-4o": 0.15, "gpt-4-turbo": 0.4, "gpt-3.5-turbo": 0.04, "gpt-4o-mini": 0.02,
    config.GEMINI_MODEL_DEFAULT: 0.2, "gemini-1.5-pro-latest": 0.2, "gemini-1.5-flash-latest": 0.05
}
KETTLE_WH: int = 360

PROTO_PATH = APP_BASE_PATH / "strategic_protocols.txt"
PROTO_TEXT: str
PROTO_HASH = ""
PROTO_LOAD_SUCCESS = False # Flag for successful load

if not PROTO_PATH.exists():
    PROTO_TEXT = config.PROTO_TEXT_FALLBACK
    logger.warning(f"Protocol file {PROTO_PATH.name} not found. Using fallback.")
    config.LOADED_PROTO_PATH_NAME = PROTO_PATH.name # For about_page.py
    config.LOADED_PROTO_TEXT = PROTO_TEXT # For about_page.py
    PROTO_LOAD_SUCCESS = False
else:
    try:
        PROTO_TEXT = PROTO_PATH.read_text(encoding="utf-8")
        PROTO_HASH = _hashlib.sha256(PROTO_TEXT.encode()).hexdigest()[:8]
        config.PROTO_TEXT_FALLBACK = PROTO_TEXT # Update fallback if successfully loaded
        config.LOADED_PROTO_PATH_NAME = PROTO_PATH.name # For about_page.py
        config.LOADED_PROTO_TEXT = PROTO_TEXT # For about_page.py
        logger.info(f"Successfully loaded protocol from {PROTO_PATH.name}")
        PROTO_LOAD_SUCCESS = True
    except Exception as e_proto:
        PROTO_TEXT = config.PROTO_TEXT_FALLBACK
        logger.error(f"Error loading protocol file {PROTO_PATH.name}: {e_proto}. Using fallback.", exc_info=True)
        config.LOADED_PROTO_PATH_NAME = PROTO_PATH.name # Still set for about_page.py
        config.LOADED_PROTO_TEXT = PROTO_TEXT # Still set for about_page.py
        PROTO_LOAD_SUCCESS = False


CH_CATEGORIES: Dict[str, str] = {
    "Accounts": "accounts", "Confirmation Stmt": "confirmation-statement", "Officers": "officers",
    "Capital": "capital", "Charges": "mortgage", "Insolvency": "insolvency",
    "PSC": "persons-with-significant-control", "Name Change": "change-of-name",
    "Reg. Office": "registered-office-address",
}

def init_session_state():
    defaults = {
        "current_topic": "general_default_topic", "session_history": [], "loaded_memories": [],
        "processed_summaries": [], # (src_id, title, summary_text) for uploaded docs/URLs
        "selected_summary_texts": [], # Texts of selected uploaded doc/URL summaries for PRIMARY context
        "latest_digest_content": "",
        "document_processing_complete": True, "ch_last_digest_path": None, "ch_last_df": None,
        "ch_last_narrative": None, "ch_last_batch_metrics": {},
        "consult_digest_model": config.OPENAI_MODEL_DEFAULT,
        "ch_analysis_summaries_for_injection": [], # List of (company_id, title_for_list, summary_text)
        
        # For "Improve Prompt" in Consult Counsel
        "user_instruction_main_text_area_value": "", # Holds current text area content
        "original_user_instruction_main": "", 
        "user_instruction_main_is_improved": False,

        # For "Improve Prompt" in CH Analysis
        "additional_ai_instructions_ch_text_area_value": "", # Holds current text area content
        "original_additional_ai_instructions_ch": "", 
        "additional_ai_instructions_ch_is_improved": False,
    }
    for key, value in defaults.items():
        if key not in st.session_state: st.session_state[key] = value
init_session_state()

with st.sidebar:
    st.markdown("## Configuration")
    current_topic_input = st.text_input("Matter / Topic ID", st.session_state.current_topic, key="topic_input_sidebar")
    if current_topic_input != st.session_state.current_topic:
        st.session_state.current_topic = current_topic_input
        # Reset topic-specific states
        st.session_state.session_history = []
        st.session_state.processed_summaries = []
        st.session_state.selected_summary_texts = []
        st.session_state.loaded_memories = []
        st.session_state.latest_digest_content = ""
        st.session_state.ch_last_digest_path = None
        st.session_state.ch_last_df = None
        st.session_state.ch_last_narrative = None
        st.session_state.ch_last_batch_metrics = {}
        st.session_state.ch_analysis_summaries_for_injection = [] # Crucially reset this
        st.rerun()

    def _topic_color_style(topic_str: str) -> str:
        color_hue = int(_hashlib.sha1(topic_str.encode()).hexdigest(), 16) % 360
        return f"background-color:hsl({color_hue}, 70%, 90%); padding:8px 12px; border-radius:8px; margin:8px 0; text-align:center; color:#333;"
    st.markdown(f"<div style='{_topic_color_style(st.session_state.current_topic)}'>Topic: <strong>{st.session_state.current_topic}</strong></div>", unsafe_allow_html=True)

    # --- Protocol Status Display ---
    st.markdown("---"); st.markdown("### System Status")
    if PROTO_LOAD_SUCCESS :
        st.success(f"Protocol '{PROTO_PATH.name}' loaded (Hash: {PROTO_HASH}).")
    elif not PROTO_PATH.exists(): # Fallback because file not found
        st.warning(f"Protocol file '{PROTO_PATH.name}' not found. Using default protocol.")
    else: # Fallback because file exists but error loading
        st.error(f"Error loading protocol '{PROTO_PATH.name}'. Using default protocol.")
    # --- End Protocol Status Display ---


    st.markdown("---"); st.markdown("### AI Model Selection")
    st.markdown("*(For Consultation & Digest Updates)*")
    if not OPENAI_API_KEY_PRESENT: st.error("‼️ OpenAI API Key missing. OpenAI models will fail.")
    if not GEMINI_API_KEY_PRESENT: st.warning("⚠️ Gemini API Key missing. Gemini models unavailable for consultation.")

    all_available_models = list(MODEL_PRICES_PER_1K_TOKENS_GBP.keys())
    gpt_models = [m for m in all_available_models if m.startswith("gpt-") and OPENAI_API_KEY_PRESENT]
    gemini_models_consult = [m for m in all_available_models if m.startswith("gemini-") and GEMINI_API_KEY_PRESENT]

    selectable_models_consult = gpt_models + gemini_models_consult
    if not selectable_models_consult: st.error("No AI models available for Consultation/Digests!");

    default_consult_model_index = 0
    if "consult_digest_model" in st.session_state and \
       st.session_state.consult_digest_model in selectable_models_consult:
        try: default_consult_model_index = selectable_models_consult.index(st.session_state.consult_digest_model)
        except ValueError: default_consult_model_index = 0 # Fallback if selected model is no longer available
    elif selectable_models_consult: # If list is not empty, set to first item
        st.session_state.consult_digest_model = selectable_models_consult[0]
    else: # No models available
        st.session_state.consult_digest_model = None


    st.session_state.consult_digest_model = st.selectbox(
        "Model for Consultation & Digests:", selectable_models_consult,
        index=default_consult_model_index,
        key="consult_digest_model_selector_main",
        disabled=not selectable_models_consult
    )
    if st.session_state.consult_digest_model:
        price_consult = MODEL_PRICES_PER_1K_TOKENS_GBP.get(st.session_state.consult_digest_model, 0.0)
        st.caption(f"Est. Cost/1k Tokens: £{price_consult:.5f}")
    else:
        st.caption("Est. Cost/1k Tokens: N/A")

    st.markdown("---")
    st.markdown("CH Summaries will use Gemini by default for speed (if configured), or fallback to OpenAI.")
    st.markdown("---")

    ai_temp = st.slider("AI Creativity (Temperature)", 0.0, 1.0, 0.2, 0.05, key="ai_temp_slider_sidebar")

    st.markdown("---"); st.markdown("### Context Injection")
    memory_file_path = APP_BASE_PATH / "memory" / f"{st.session_state.current_topic}.json"
    loaded_memories_from_file: List[str] = []
    if memory_file_path.exists():
        try:
            mem_data = json.loads(memory_file_path.read_text(encoding="utf-8"))
            if isinstance(mem_data, list):
                loaded_memories_from_file = [str(item) for item in mem_data if isinstance(item, str)]
        except Exception as e_mem_load: st.warning(f"Could not load memory file {memory_file_path.name}: {e_mem_load}")
    selected_mem_snippets = st.multiselect("Inject Memories", loaded_memories_from_file,
        default=[mem for mem in st.session_state.loaded_memories if mem in loaded_memories_from_file], # Persist selection
        key="mem_multiselect_sidebar")
    st.session_state.loaded_memories = selected_mem_snippets

    digest_file_path = APP_BASE_PATH / "memory" / "digests" / f"{st.session_state.current_topic}.md"
    if digest_file_path.exists():
        try: st.session_state.latest_digest_content = digest_file_path.read_text(encoding="utf-8")
        except Exception as e_digest_load: st.warning(f"Could not load digest {digest_file_path.name}: {e_digest_load}"); st.session_state.latest_digest_content = ""
    else: st.session_state.latest_digest_content = "" # Ensure it's empty if file doesn't exist
    inject_digest_checkbox = st.checkbox("Inject Digest", value=bool(st.session_state.latest_digest_content), # Default based on content
        key="inject_digest_checkbox_sidebar", disabled=not bool(st.session_state.latest_digest_content))

    st.markdown("---"); st.markdown("### Document Intake (for Context)")
    uploaded_docs_list = st.file_uploader("Upload Docs (PDF, DOCX, TXT)", ["pdf", "docx", "txt"],
        accept_multiple_files=True, key="doc_uploader_sidebar")
    urls_input_str = st.text_area("Paste URLs (one per line)", key="url_textarea_sidebar", height=80)
    urls_to_process = [u.strip() for u in urls_input_str.splitlines() if u.strip().startswith("http")]

    current_source_identifiers = {f.name for f in uploaded_docs_list} | set(urls_to_process)
    processed_summary_ids_in_session = {s_tuple[0] for s_tuple in st.session_state.processed_summaries}
    sources_needing_processing = current_source_identifiers - processed_summary_ids_in_session

    newly_processed_summaries_for_this_run_sidebar: List[Tuple[str, str, str]] = [] # Define here for wider scope
    if sources_needing_processing and st.session_state.document_processing_complete:
        st.session_state.document_processing_complete = False # Prevent re-processing during rerun
        summaries_cache_dir_for_topic = APP_BASE_PATH / "summaries" / st.session_state.current_topic
        summaries_cache_dir_for_topic.mkdir(parents=True, exist_ok=True)

        with st.spinner(f"Processing {len(sources_needing_processing)} new document(s)/URL(s)..."):
            progress_bar_docs = st.progress(0.0)
            for idx, src_id in enumerate(list(sources_needing_processing)): # Convert set to list for indexing
                title, summary = "Error", "Processing Failed"
                # Simple cache key based on source identifier hash
                cache_file_name = f"summary_{_hashlib.sha256(src_id.encode()).hexdigest()[:16]}.json"
                summary_cache_file = summaries_cache_dir_for_topic / cache_file_name

                if summary_cache_file.exists():
                    try:
                        cached_data = json.loads(summary_cache_file.read_text(encoding="utf-8"))
                        title, summary = cached_data.get("t", "Cache Title Error"), cached_data.get("s", "Cache Summary Error")
                    except Exception: title, summary = "Error", "Processing Failed (Cache Read)" # More specific cache error

                if title == "Error" or "Cache" in title : # If cache load failed or it was an error state
                    raw_content, error_msg = None, None
                    # Check if it's an uploaded file or a URL
                    if src_id in {f.name for f in uploaded_docs_list}: # Is it an uploaded file?
                        file_obj = next((f for f in uploaded_docs_list if f.name == src_id), None)
                        if file_obj: raw_content, error_msg = extract_text_from_uploaded_file(io.BytesIO(file_obj.getvalue()), src_id)
                    elif src_id in urls_to_process: # Is it a URL?
                        raw_content, error_msg = fetch_url_content(src_id)

                    if error_msg: title, summary = f"Error: {src_id[:40]}...", error_msg
                    elif not raw_content or not raw_content.strip(): title, summary = f"Empty: {src_id[:40]}...", "No text content found or extracted."
                    else: # Successfully got raw content
                        # Use a cost-effective model for these quick summaries
                        title, summary = summarise_with_title(raw_content, "gpt-4o-mini", st.session_state.current_topic, PROTO_TEXT)

                    if "Error" not in title and "Empty" not in title: # Cache if successfully processed
                        try: summary_cache_file.write_text(json.dumps({"t":title,"s":summary,"src":src_id}),encoding="utf-8")
                        except Exception as e_c: logger.warning(f"Cache write failed for {src_id}: {e_c}")

                newly_processed_summaries_for_this_run_sidebar.append((src_id, title, summary))
                progress_bar_docs.progress((idx + 1) / len(sources_needing_processing))

            # Update session state: keep existing ones that are still valid, add new ones
            existing_to_keep = [s for s in st.session_state.processed_summaries if s[0] in current_source_identifiers and s[0] not in sources_needing_processing]
            st.session_state.processed_summaries = existing_to_keep + newly_processed_summaries_for_this_run_sidebar
            progress_bar_docs.empty()
        st.session_state.document_processing_complete = True; st.rerun() # Rerun to update UI with new summaries

    # Selection for Uploaded/URL Summaries
    st.session_state.selected_summary_texts = [] # Reset before populating based on checkbox state
    if st.session_state.processed_summaries:
        st.markdown("---"); st.markdown("### Available Doc/URL Summaries (Select to Inject)")
        for idx, (s_id, title, summary_text) in enumerate(st.session_state.processed_summaries):
            checkbox_key = f"sum_sel_{_hashlib.md5(s_id.encode()).hexdigest()}"
            is_newly_processed = any(s_id == item[0] for item in newly_processed_summaries_for_this_run_sidebar)
            # Default to checked if newly processed, or if previously checked (and still exists)
            default_checked = is_newly_processed or st.session_state.get(checkbox_key, False)
            is_injected = st.checkbox(f"{idx+1}. {title[:40]}...", value=default_checked, key=checkbox_key, help=f"Source: {s_id}\nSummary: {summary_text[:200]}...")
            if is_injected: st.session_state.selected_summary_texts.append(f"UPLOADED DOCUMENT/URL SUMMARY ('{title}'):\n{summary_text}")


    # Selection for CH Analysis Summaries
    selected_ch_summary_texts_for_injection_temp = [] # Temp list for this run
    if st.session_state.ch_analysis_summaries_for_injection:
        st.markdown("---"); st.markdown("### CH Analysis Summaries (Select to Inject)")
        for idx, (company_id, title_for_list, summary_text) in enumerate(st.session_state.ch_analysis_summaries_for_injection):
            ch_checkbox_key = f"ch_sum_sel_{_hashlib.md5(company_id.encode() + title_for_list.encode()).hexdigest()}"
            # Default to False unless explicitly checked by the user.
            is_ch_summary_injected = st.checkbox(f"{idx+1}. CH: {title_for_list[:40]}...", value=st.session_state.get(ch_checkbox_key, False), key=ch_checkbox_key, help=f"Company: {company_id}\nSummary: {summary_text[:200]}...")
            if is_ch_summary_injected:
                selected_ch_summary_texts_for_injection_temp.append(f"COMPANIES HOUSE ANALYSIS SUMMARY FOR {company_id} ({title_for_list}):\n{summary_text}")
    # This state is now dynamically built when creating context for AI rather than storing selection list in session_state permanently


    st.markdown("---")
    if st.button("End Session & Update Digest", key="end_session_button_sidebar"):
        if not st.session_state.session_history: st.warning("No new interactions to add to digest.")
        elif not st.session_state.consult_digest_model: st.error("No AI model selected for Digest Update.")
        else:
            with st.spinner("Updating Digest..."):
                new_interactions_block = "\n\n---\n\n".join(st.session_state.session_history)
                existing_digest_text = st.session_state.latest_digest_content
                update_digest_prompt = (f"Consolidate the following notes. Integrate the NEW interactions into the EXISTING digest, "
                                    f"maintaining a coherent and concise summary. Aim for a maximum of around 2000 words for the entire updated digest. "
                                    f"Preserve key facts and decisions.\n\n"
                                    f"EXISTING DIGEST (for topic: {st.session_state.current_topic}):\n{existing_digest_text}\n\n"
                                    f"NEW INTERACTIONS (to integrate for topic: {st.session_state.current_topic}):\n{new_interactions_block}")
                try:
                    current_ai_model_for_digest = st.session_state.consult_digest_model
                    updated_digest_text = "Error updating digest."
                    if current_ai_model_for_digest.startswith("gpt-"):
                        client = config.get_openai_client(); assert client
                        resp = client.chat.completions.create(model=current_ai_model_for_digest, temperature=0.1, max_tokens=3000, messages=[{"role": "system", "content": PROTO_TEXT}, {"role": "user", "content": update_digest_prompt}])
                        updated_digest_text = resp.choices[0].message.content.strip()
                    elif current_ai_model_for_digest.startswith("gemini-"):
                        client = config.get_gemini_model(current_ai_model_for_digest); assert client and config.genai # Check config.genai
                        full_prompt_gemini = f"{PROTO_TEXT}\n\n{update_digest_prompt}" # Combine for Gemini
                        resp = client.generate_content(full_prompt_gemini, generation_config=config.genai.types.GenerationConfig(temperature=0.1, max_output_tokens=3000)) # Use config.genai
                        updated_digest_text = resp.text.strip() # Check for block reason

                    digest_file_path.write_text(updated_digest_text, encoding="utf-8")
                    historical_digest_path = APP_BASE_PATH / "memory" / "digests" / f"history_{st.session_state.current_topic}.md"
                    with historical_digest_path.open("a", encoding="utf-8") as fp_hist:
                        fp_hist.write(f"\n\n### Update: {_dt.datetime.now():%Y-%m-%d %H:%M} (Model: {current_ai_model_for_digest})\n{updated_digest_text}\n---\n")
                    st.success(f"Digest for '{st.session_state.current_topic}' updated."); st.session_state.session_history = []; st.session_state.latest_digest_content = updated_digest_text; st.rerun()
                except Exception as e_digest_update:
                    st.error(f"Digest update failed: {e_digest_update}"); logger.error(f"Digest update error: {e_digest_update}", exc_info=True)

# ── Main Application Area UI (Using Tabs) ─────────────────────────
st.markdown(f"## 🏛️ Strategic Counsel: {st.session_state.current_topic}")
tab_consult, tab_ch_analysis, tab_about_rendered = st.tabs(["💬 Consult Counsel", "🇬🇧 Companies House Analysis", "ℹ️ About"])

with tab_consult:
    st.markdown("Provide instructions and context (using sidebar options) for drafting, analysis, or advice.")
    
    # Text area for user's main instruction. Its value is stored in st.session_state.main_instruction_area_consult_tab (by Streamlit)
    # and mirrored to st.session_state.user_instruction_main_text_area_value by the on_change callback.
    st.text_area(
        "Your Instruction:", 
        value=st.session_state.user_instruction_main_text_area_value, # Display value from our dedicated session state
        height=200, 
        key="main_instruction_area_consult_tab", # Key for this specific widget
        on_change=lambda: st.session_state.update(user_instruction_main_text_area_value=st.session_state.main_instruction_area_consult_tab) # Update our dedicated state from widget's state
    )

    col_improve_main, col_cancel_main, col_spacer_main = st.columns([2,2,3]) # Adjusted column ratios
    with col_improve_main:
        if st.button("💡 Suggest Improved Prompt", key="improve_prompt_main_button", help="Let AI refine your instruction for better results.", use_container_width=True):
            current_text_in_area = st.session_state.user_instruction_main_text_area_value 
            if current_text_in_area and current_text_in_area.strip():
                if not st.session_state.user_instruction_main_is_improved: 
                    st.session_state.original_user_instruction_main = current_text_in_area
                
                with st.spinner("Improving prompt..."):
                    improved_prompt = get_improved_prompt(current_text_in_area, "Strategic Counsel general query")
                    if "Error:" not in improved_prompt and improved_prompt.strip():
                        st.session_state.user_instruction_main_text_area_value = improved_prompt 
                        st.session_state.user_instruction_main_is_improved = True
                        st.rerun() 
                    elif "Error:" in improved_prompt:
                        st.warning(f"Could not improve prompt: {improved_prompt}")
                    # If prompt is empty or only whitespace after improvement, no change is made to the text area.
            else:
                st.info("Please enter an instruction first to improve it.")

    with col_cancel_main:
        if st.session_state.user_instruction_main_is_improved:
            if st.button("↩️ Revert to Original", key="cancel_improve_prompt_main_button", use_container_width=True):
                st.session_state.user_instruction_main_text_area_value = st.session_state.original_user_instruction_main
                st.session_state.user_instruction_main_is_improved = False
                st.rerun()

    consult_model_name = st.session_state.get("consult_digest_model")

    if st.button("✨ Consult Counsel", type="primary", key="run_ai_button_consult_tab"):
        current_instruction_to_use = st.session_state.user_instruction_main_text_area_value

        if not current_instruction_to_use.strip(): st.warning("Please enter your instructions.")
        elif not consult_model_name: st.error("No AI model selected for Consultation.")
        else:
            with st.spinner(f"Consulting {consult_model_name}..."):
                messages_for_ai = [{"role": "system", "content": PROTO_TEXT + f"\n[Protocol Hash:{PROTO_HASH}]"}]
                context_parts_for_ai = []
                if inject_digest_checkbox and st.session_state.latest_digest_content: context_parts_for_ai.append(f"CURRENT DIGEST:\n{st.session_state.latest_digest_content}")
                if st.session_state.loaded_memories: context_parts_for_ai.append("INJECTED MEMORIES:\n" + "\n---\n".join(st.session_state.loaded_memories))

                combined_selected_summaries = []
                if st.session_state.selected_summary_texts: 
                    combined_selected_summaries.extend(st.session_state.selected_summary_texts)
                
                if "ch_analysis_summaries_for_injection" in st.session_state and st.session_state.ch_analysis_summaries_for_injection:
                    for idx, (company_id, title_for_list, summary_text) in enumerate(st.session_state.ch_analysis_summaries_for_injection):
                        ch_checkbox_key = f"ch_sum_sel_{_hashlib.md5(company_id.encode() + title_for_list.encode()).hexdigest()}"
                        if st.session_state.get(ch_checkbox_key, False): 
                            combined_selected_summaries.append(f"COMPANIES HOUSE ANALYSIS SUMMARY FOR {company_id} ({title_for_list}):\n{summary_text}")
                
                if combined_selected_summaries:
                    context_parts_for_ai.append("SELECTED DOCUMENT SUMMARIES & ANALYSIS:\n" + "\n===\n".join(combined_selected_summaries))

                if context_parts_for_ai: messages_for_ai.append({"role": "system", "content": "ADDITIONAL CONTEXT:\n\n" + "\n\n".join(context_parts_for_ai)})
                messages_for_ai.append({"role": "user", "content": current_instruction_to_use}) # Use the potentially improved instruction

                try:
                    ai_response_text = "Error: AI response could not be generated."
                    prompt_tokens, completion_tokens = 0, 0

                    if consult_model_name.startswith("gpt-"):
                        openai_client = config.get_openai_client(); assert openai_client
                        ai_api_response = openai_client.chat.completions.create(
                            model=consult_model_name, temperature=ai_temp, messages=messages_for_ai, max_tokens=3500
                        )
                        ai_response_text = ai_api_response.choices[0].message.content.strip()
                        if ai_api_response.usage:
                            prompt_tokens = ai_api_response.usage.prompt_tokens
                            completion_tokens = ai_api_response.usage.completion_tokens
                    elif consult_model_name.startswith("gemini-"):
                        gemini_model_client = config.get_gemini_model(consult_model_name); assert gemini_model_client and config.genai
                        try: 
                            full_prompt_str_gemini = "\n\n".join([f"{m['role']}: {m['content']}" for m in messages_for_ai])
                            count_resp_prompt = gemini_model_client.count_tokens(full_prompt_str_gemini)
                            prompt_tokens = count_resp_prompt.total_tokens
                        except Exception as e_gem_count_p: logger.warning(f"Gemini prompt token count failed: {e_gem_count_p}"); prompt_tokens = 0

                        gemini_api_response = gemini_model_client.generate_content(
                            contents=messages_for_ai,
                            generation_config=config.genai.types.GenerationConfig(temperature=ai_temp, max_output_tokens=3500)
                        )
                        if hasattr(gemini_api_response, 'text') and gemini_api_response.text:
                             ai_response_text = gemini_api_response.text.strip()
                             try: 
                                 count_resp_completion = gemini_model_client.count_tokens(ai_response_text)
                                 completion_tokens = count_resp_completion.total_tokens
                             except Exception as e_gem_count_c: logger.warning(f"Gemini completion token count failed: {e_gem_count_c}"); completion_tokens = 0
                        elif hasattr(gemini_api_response, 'prompt_feedback') and gemini_api_response.prompt_feedback.block_reason:
                            block_reason_str = config.genai.types.BlockedReason(gemini_api_response.prompt_feedback.block_reason).name
                            ai_response_text = f"Error: Gemini content generation blocked. Reason: {block_reason_str}."
                            logger.error(f"Gemini content blocked. Reason: {block_reason_str}. Feedback: {gemini_api_response.prompt_feedback}")
                        else:
                             ai_response_text = "Error: Gemini response was empty or malformed."
                             logger.error(f"Gemini empty/malformed response: {gemini_api_response}")
                    else:
                        raise ValueError(f"Unsupported model type for consultation: {consult_model_name}")

                    st.session_state.session_history.append(f"Instruction:\n{current_instruction_to_use}\n\nResponse ({consult_model_name}):\n{ai_response_text}") # Log the used instruction
                    with st.chat_message("assistant", avatar="⚖️"): st.markdown(ai_response_text)

                    with st.expander("📊 Run Details & Export"):
                        total_tokens = prompt_tokens + completion_tokens
                        cost = (total_tokens / 1000) * MODEL_PRICES_PER_1K_TOKENS_GBP.get(consult_model_name,0.0) if total_tokens > 0 else 0.0
                        energy_model_wh = MODEL_ENERGY_WH_PER_1K_TOKENS.get(consult_model_name, 0.0)
                        energy_wh = (total_tokens / 1000) * energy_model_wh if total_tokens > 0 else 0.0

                        st.metric("Total Tokens", f"{total_tokens:,}", f"{prompt_tokens:,} prompt + {completion_tokens:,} completion")
                        st.metric("Est. Cost", f"£{cost:.5f}")
                        if energy_model_wh > 0 and energy_wh > 0:
                            st.metric("Est. Energy", f"{energy_wh:.3f}Wh", f"~{(energy_wh / KETTLE_WH * 100):.1f}% Kettle" if KETTLE_WH > 0 else "")

                        ts_now_str = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
                        docx_filename = f"{st.session_state.current_topic}_{ts_now_str}_response.docx"
                        docx_export_path = APP_BASE_PATH / "exports" / docx_filename
                        try:
                            doc = Document(); doc.add_heading(f"AI Consultation: {st.session_state.current_topic}",0)
                            doc.add_paragraph(f"Instruction:\n{current_instruction_to_use}\n\nResponse ({consult_model_name} @ {ts_now_str}):\n{ai_response_text}") # Use current_instruction_to_use
                            doc.save(docx_export_path)
                            with open(docx_export_path, "rb") as fp_docx: st.download_button("Download .docx", fp_docx, docx_filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        except Exception as e_docx: st.error(f"DOCX export error: {e_docx}")

                        log_filename = f"{st.session_state.current_topic}_{ts_now_str}_log.json"
                        log_export_path = APP_BASE_PATH / "logs" / log_filename
                        log_data = {"topic":st.session_state.current_topic, "timestamp":ts_now_str, "model":consult_model_name, "temp":ai_temp, "tokens":{"p":prompt_tokens,"c":completion_tokens,"t":total_tokens}, "cost_gbp":cost, "energy_wh":energy_wh, "user_instr":current_instruction_to_use[:200]+("..." if len(current_instruction_to_use) > 200 else ""), "resp_preview":ai_response_text[:200]+("..." if len(ai_response_text) > 200 else "")} # Use current_instruction_to_use
                        try: log_export_path.write_text(json.dumps(log_data, indent=2), encoding="utf-8")
                        except Exception as e_log: st.error(f"Log save error: {e_log}")

                except Exception as e_ai_consult:
                    st.error(f"AI Consultation Error with {consult_model_name}: {e_ai_consult}", icon="🚨")
                    logger.error(f"AI Consultation Error ({consult_model_name}): {e_ai_consult}", exc_info=True)

    if st.session_state.session_history:
        st.markdown("---"); st.subheader("📜 Current Session History (Newest First)")
        history_display_container = st.container(height=400) # Ensure fixed height for scroll
        for i, entry_text in enumerate(reversed(st.session_state.session_history)):
            history_display_container.markdown(f"**Interaction {len(st.session_state.session_history)-i}:**\n---\n{entry_text}\n\n")

with tab_ch_analysis:
    st.markdown("Fetch, process, and summarize UK company filings. Summaries will use Gemini (if configured) for optimal speed and context, or fallback to OpenAI.")
    if not CH_API_KEY_PRESENT: st.warning("⚠️ CH API Key missing. This tab will be limited.", icon="🔑")

    col1_ch_input, col2_ch_results = st.columns([1, 2])
    with col1_ch_input:
        st.subheader("Inputs & Configuration")
        ch_input_method = st.radio("1. Input Method:", ("Upload CSV", "Single Lookup"), key="ch_input_method_radio", horizontal=True, disabled=not CH_API_KEY_PRESENT)
        uploaded_ch_csv_file_ui = None; single_company_query_str_ui = ""
        if ch_input_method == "Upload CSV": uploaded_ch_csv_file_ui = st.file_uploader("Upload CSV", ["csv"], key="ch_csv_uploader_ui")
        else: single_company_query_str_ui = st.text_input("Company Name or Number", key="ch_single_company_input_ui")

        st.markdown("---"); st.markdown("##### Document Selection")
        selected_ch_categories_display = st.multiselect("Categories", list(CH_CATEGORIES.keys()), default=["Accounts", "Confirmation Stmt", "Charges"], key="ch_categories_multiselect_ui")
        api_categories_for_pipeline = [CH_CATEGORIES[cat_name] for cat_name in selected_ch_categories_display if cat_name in CH_CATEGORIES]
        current_system_year = _dt.date.today().year
        ch_default_start_year, ch_default_end_year = current_system_year - 4, current_system_year - 1 # Default to last 3 full years
        selected_year_range_ch = st.slider("Filing Year Range", 2000, current_system_year, (ch_default_start_year, ch_default_end_year), key="ch_year_range_slider_ui")
        start_year_for_pipeline, end_year_for_pipeline = selected_year_range_ch

        st.markdown("---"); st.markdown("##### Analysis & Output Options")
        
        st.text_area(
            "Additional AI Summary Instructions", 
            value=st.session_state.additional_ai_instructions_ch_text_area_value,
            placeholder="Example: Focus on director changes, dividend policy...", 
            key="ch_ai_instructions_textarea_ui", 
            height=100,
            on_change=lambda: st.session_state.update(additional_ai_instructions_ch_text_area_value=st.session_state.ch_ai_instructions_textarea_ui)
        )

        col_improve_ch, col_cancel_ch, col_spacer_ch = st.columns([2,2,3]) # Adjusted column ratios
        with col_improve_ch:
            if st.button("💡 Suggest Improved Instruction", key="improve_prompt_ch_button", help="Let AI refine your CH analysis instruction.", use_container_width=True):
                current_text_in_ch_area = st.session_state.additional_ai_instructions_ch_text_area_value
                if current_text_in_ch_area and current_text_in_ch_area.strip():
                    if not st.session_state.additional_ai_instructions_ch_is_improved:
                        st.session_state.original_additional_ai_instructions_ch = current_text_in_ch_area
                    
                    with st.spinner("Improving instruction..."):
                        improved_ch_instruction = get_improved_prompt(current_text_in_ch_area, "Companies House document analysis instruction")
                        if "Error:" not in improved_ch_instruction and improved_ch_instruction.strip():
                            st.session_state.additional_ai_instructions_ch_text_area_value = improved_ch_instruction
                            st.session_state.additional_ai_instructions_ch_is_improved = True
                            st.rerun()
                        elif "Error:" in improved_ch_instruction:
                            st.warning(f"Could not improve CH instruction: {improved_ch_instruction}")
                else:
                    st.info("Please enter CH analysis instructions first to improve them.")
        
        with col_cancel_ch:
            if st.session_state.additional_ai_instructions_ch_is_improved:
                if st.button("↩️ Revert to Original", key="cancel_improve_prompt_ch_button", use_container_width=True):
                    st.session_state.additional_ai_instructions_ch_text_area_value = st.session_state.original_additional_ai_instructions_ch
                    st.session_state.additional_ai_instructions_ch_is_improved = False
                    st.rerun()

        st.markdown("---"); st.markdown("##### Content Focusing (Optional)")
        ch_keywords_for_filtering = st.text_input(
            "Keywords/Topics for Focused Summary (comma-separated)",
            placeholder="e.g., net profit, director remuneration, sustainability",
            key="ch_keywords_filtering_input_ui",
            help="If provided, the AI prompt will be guided to focus on these. Full text pre-filtering is a future enhancement."
        )

        use_textract_ocr_ch = st.checkbox("Use AWS Textract for PDF OCR", value=False, key="ch_use_textract_checkbox_ui", help="Requires AWS setup.", disabled=not CH_PIPELINE_TEXTRACT_FLAG_FROM_MODULE)
        if not CH_PIPELINE_TEXTRACT_FLAG_FROM_MODULE and use_textract_ocr_ch: st.warning("Textract OCR selected, but backend utilities unavailable."); use_textract_ocr_ch = False # Auto-disable if not available
        keep_temp_files_days_ch = st.slider("Temp File Retention (Days)", 0, 30, 1, key="ch_keep_temp_files_slider_ui")

        max_docs_to_fetch_meta = st.sidebar.number_input("Max initial documents to scan per category (CH meta)", min_value=10, max_value=500, value=100, step=10)
        target_docs_per_category_in_date_range = st.sidebar.number_input("Target documents per category (within date range)", min_value=1, max_value=100, value=10, step=1) # New input

        run_ch_analysis_button_ui = st.button("🚀 Run Companies House Analysis", key="ch_run_analysis_button_ui", use_container_width=True, disabled=not CH_API_KEY_PRESENT)

    with col2_ch_results:
        st.subheader("Analysis Results")
        ch_results_display_container = st.container(border=True, height=600) # Keep a container for scrollability

        if run_ch_analysis_button_ui:
            st.session_state.ch_last_digest_path = None; st.session_state.ch_last_df = None
            st.session_state.ch_last_narrative = None; st.session_state.ch_last_batch_metrics = {}
            st.session_state.ch_analysis_summaries_for_injection = [] # Clear previous CH summaries for injection list

            with ch_results_display_container: # Keep spinner and initial messages inside this container
                st.info("Preparing CH analysis run...")
            # ... (rest of the CSV preparation logic remains the same) ...
                temp_csv_for_pipeline_path: Optional[_pl.Path] = None
                # Use topic-specific scratch directory to avoid conflicts if multiple users/topics run simultaneously
                # This also makes cleanup per topic easier if desired, though global keep_days still applies
                run_specific_scratch_base = APP_BASE_PATH / "temp_ch_runs" / st.session_state.current_topic
                run_specific_scratch_base.mkdir(parents=True, exist_ok=True)
                ch_run_scratch_dir = _pl.Path(tempfile.mkdtemp(prefix=f"ch_run_{_dt.datetime.now():%Y%m%d%H%M%S}_", dir=run_specific_scratch_base))

                try: ch_run_scratch_dir.mkdir(parents=True, exist_ok=True) # Ensure it exists, mkdtemp should create it
                except OSError as e_mkdir_ch_run:
                    with ch_results_display_container: st.error(f"Error creating run scratch dir: {e_mkdir_ch_run}");
                    st.stop() # Stop execution if scratch dir fails

                # CSV preparation logic (Upload or Single Lookup)
                if ch_input_method == "Upload CSV":
                    if not uploaded_ch_csv_file_ui:
                        with ch_results_display_container: st.warning("Please upload CSV."); st.stop()
                    temp_csv_for_pipeline_path = ch_run_scratch_dir / uploaded_ch_csv_file_ui.name
                    try: temp_csv_for_pipeline_path.write_bytes(uploaded_ch_csv_file_ui.getbuffer())
                    except IOError as e_csv_save:
                        with ch_results_display_container: st.error(f"Error saving CSV: {e_csv_save}"); st.stop()
                else: # Single Lookup
                    if not single_company_query_str_ui.strip():
                        with ch_results_display_container: st.warning("Enter company name/number."); st.stop()
                    found_co_no, find_err, _ = find_company_number(single_company_query_str_ui, config.CH_API_KEY)
                    if find_err:
                        with ch_results_display_container: st.error(f"Lookup Error: {find_err}"); st.stop()
                    if not found_co_no:
                        with ch_results_display_container: st.error("Could not resolve company."); st.stop()
                    with ch_results_display_container: st.success(f"Found: Using company number {found_co_no}.")
                    temp_csv_for_pipeline_path = ch_run_scratch_dir / f"single_co_{found_co_no}.csv"
                    try:
                        with open(temp_csv_for_pipeline_path, "w", newline="", encoding="utf-8") as f_csv:
                            writer = csv.writer(f_csv); writer.writerow(["CompanyNumber"]); writer.writerow([found_co_no])
                    except IOError as e_temp_csv:
                        with ch_results_display_container: st.error(f"Error creating temp CSV: {e_temp_csv}"); st.stop()

                if not temp_csv_for_pipeline_path or not temp_csv_for_pipeline_path.exists():
                    with ch_results_display_container: st.error("Input file for pipeline not prepared."); st.stop()
                if not api_categories_for_pipeline:
                    with ch_results_display_container: st.warning("Select document categories."); st.stop()
            # End of input prep, spinner for the main processing
            with st.spinner(f"Running CH analysis (Summaries via Gemini default, Textract: {'On' if use_textract_ocr_ch else 'Off'})..."):
                current_ch_instruction_to_use = st.session_state.additional_ai_instructions_ch_text_area_value
                output_digest_file_path, batch_metrics_from_run = run_batch_company_analysis(
                    csv_path=temp_csv_for_pipeline_path, selected_categories=api_categories_for_pipeline,
                    start_year=start_year_for_pipeline, end_year=end_year_for_pipeline,
                    model_prices_gbp=MODEL_PRICES_PER_1K_TOKENS_GBP,
                    specific_ai_instructions=current_ch_instruction_to_use, # Pass the used instruction
                    filter_keywords_str=ch_keywords_for_filtering, 
                    base_scratch_dir=ch_run_scratch_dir.parent, 
                    keep_days=keep_temp_files_days_ch,
                    max_docs_to_fetch_meta=max_docs_to_fetch_meta, # Pass new parameter
                    target_docs_per_category_in_date_range=target_docs_per_category_in_date_range # Pass new parameter
                    # use_textract_ocr is passed to ch_pipeline
                )
                st.session_state.ch_last_batch_metrics = batch_metrics_from_run
                if output_digest_file_path and output_digest_file_path.exists():
                    st.session_state.ch_last_digest_path = output_digest_file_path
                    try:
                        st.session_state.ch_last_df = pd.read_csv(output_digest_file_path)
                        # Populate summaries for injection from the newly loaded DataFrame
                        if st.session_state.ch_last_df is not None and not st.session_state.ch_last_df.empty:
                            temp_ch_summaries_for_injection = []
                            required_cols_for_injection = ['parent_company_no', 'summary_of_findings']
                            if all(col in st.session_state.ch_last_df.columns for col in required_cols_for_injection):
                                for _, row_data in st.session_state.ch_last_df.iterrows():
                                    company_id_val = str(row_data.get('parent_company_no', 'UnknownCompany'))
                                    summary_text_val = str(row_data.get('summary_of_findings', ''))
                                    # Add only if summary is meaningful
                                    if summary_text_val and "Error:" not in summary_text_val and "No content provided" not in summary_text_val and len(summary_text_val) > 50 :
                                        list_title = f"Co. {company_id_val}" # Simple title for the selection list
                                        temp_ch_summaries_for_injection.append((company_id_val, list_title, summary_text_val))
                                st.session_state.ch_analysis_summaries_for_injection = temp_ch_summaries_for_injection
                                logger.info(f"Populated {len(temp_ch_summaries_for_injection)} CH summaries for potential injection into Counsel context.")
                            else:
                                logger.warning("CH results digest CSV is missing 'parent_company_no' or 'summary_of_findings', cannot populate for injection.")
                    except Exception as e_read_csv:
                        with ch_results_display_container: st.error(f"Error reading digest CSV: {e_read_csv}");
                        st.session_state.ch_last_df = None # Ensure df is None on error
                elif batch_metrics_from_run.get("error"): # If pipeline returned an error in metrics
                    with ch_results_display_container: st.error(f"CH Pipeline Error: {batch_metrics_from_run['error']}");
                    st.session_state.ch_last_df = None
                else: # No error, but no output path or empty DataFrame from CSV
                    with ch_results_display_container:
                        if "No companies processed" in batch_metrics_from_run.get("notes",""): st.warning("CH Analysis: No companies processed.")
                        else: st.error("CH Analysis ran, but output CSV not found or was empty.")
            st.rerun() # Rerun to update sidebar with new CH summaries & display results

        # Display logic for CH results, always using the ch_results_display_container
        with ch_results_display_container:
            if st.session_state.ch_last_df is not None and not st.session_state.ch_last_df.empty:
                st.success("✅ Companies House Batch Analysis Complete!")
                st.markdown("---")

                if st.session_state.ch_last_digest_path and _pl.Path(st.session_state.ch_last_digest_path).exists():
                    try:
                        with open(st.session_state.ch_last_digest_path, "rb") as fp_dl: # Read as bytes
                            st.download_button(
                                label="Download Digest CSV",
                                data=fp_dl, # Pass bytes data
                                file_name=_pl.Path(st.session_state.ch_last_digest_path).name,
                                mime="text/csv",
                                key="ch_dl_digest_btn_main_view_updated"
                            )
                    except Exception as e_dl:
                        st.warning(f"Download button error: {e_dl}")

                if st.session_state.ch_last_batch_metrics:
                    with st.expander("Batch Processing Metrics (Last Run)", expanded=False):
                        metrics = st.session_state.ch_last_batch_metrics
                        st.text(f"Run Timestamp: {metrics.get('run_timestamp', 'N/A')}")
                        st.text(f"Parent Companies Processed: {metrics.get('total_parent_companies_processed', 0)}")
                        st.text(f"Successfully Summarized: {metrics.get('companies_successfully_summarized', 0)}")
                        st.text(f"Keywords Applied to Batch: {metrics.get('keywords_applied_to_batch', 'N/A')}")


                        ai_costs = metrics.get("ai_ch_summary_costs", {})
                        if ai_costs and ai_costs.get("model_used_for_ch_summaries"):
                            st.markdown("###### AI Summarization Costs (CH Documents):")
                            st.text(f"  Model Used: {ai_costs.get('model_used_for_ch_summaries')}")
                            st.text(f"  Total Prompt Tokens: {ai_costs.get('total_prompt_tokens', 0):,}")
                            st.text(f"  Total Completion Tokens: {ai_costs.get('total_completion_tokens', 0):,}")
                            st.text(f"  Estimated Cost: £{ai_costs.get('estimated_cost_gbp', 0.0):.5f}")

                        aws_costs = metrics.get("aws_ocr_costs", {})
                        if aws_costs and any(val for key, val in aws_costs.items() if key != "notes" and val !=0): # Check for actual cost/usage
                            st.markdown("###### AWS Textract OCR Cost Estimation (if used):")
                            st.text(f"  Textract Pages Processed: {metrics.get('total_textract_pages_processed',0)}")
                            st.text(f"  PDFs Sent to Textract: {metrics.get('total_pdfs_sent_to_textract',0)}")
                            for key, value in aws_costs.items():
                                if key != "notes": st.text(f"  {key.replace('_', ' ').title()}: {value}")
                            if "notes" in aws_costs and aws_costs["notes"]: st.caption(f"  Cost Notes: {aws_costs['notes']}")
                        elif "notes" in aws_costs and aws_costs["notes"]: st.caption(f"OCR Cost Notes: {aws_costs['notes']}")
                st.markdown("---")
                st.markdown("##### Companies House Analysis Results")
                required_cols_display = ['parent_company_no', 'summary_of_findings']
                if all(col in st.session_state.ch_last_df.columns for col in required_cols_display):
                    for index, row in st.session_state.ch_last_df.iterrows():
                        company_identifier = str(row['parent_company_no'])
                        summary_content = str(row.get('summary_of_findings','')) # Ensure it's a string
                        char_count_val = row.get('combined_text_char_count', 0)
                        # Ensure char_count_val is numeric before formatting
                        char_count_str = ""
                        if isinstance(char_count_val, (int, float)) and char_count_val > 0:
                            char_count_str = f"{int(char_count_val):,}"
                        elif isinstance(char_count_val, str) and char_count_val.isdigit() and int(char_count_val) > 0:
                             char_count_str = f"{int(char_count_val):,}"


                        expander_title = f"Company: {company_identifier}"
                        if char_count_str:
                            expander_title += f" (Summary from ~{char_count_str} chars of source text)"

                        with st.expander(expander_title):
                            if summary_content and summary_content.strip() and "Error:" not in summary_content :
                                # Attempt to highlight Red Flags section
                                summary_lower_case = summary_content.lower()
                                red_flag_heading_keyword = "red flags / key risks:" # Keyword to search for
                                red_flag_section_pos = summary_lower_case.find(red_flag_heading_keyword)

                                if red_flag_section_pos != -1:
                                    text_before_flags = summary_content[:red_flag_section_pos]
                                    # Find the actual start of the heading in original case for display
                                    actual_heading_start_index = summary_lower_case.find(red_flag_heading_keyword)
                                    # Extract the heading and the content after it
                                    red_flags_full_section = summary_content[actual_heading_start_index:]
                                    
                                    # Split the red flags section from any subsequent standard sections if AI uses them
                                    match_next_std_section = re.search(r"\n\s*\d+\.\s+[A-Z]{2,}", red_flags_full_section[len(red_flag_heading_keyword):], re.MULTILINE)
                                    
                                    actual_red_flag_heading_text = red_flags_full_section[:len(red_flag_heading_keyword)]
                                    red_flag_content_only = red_flags_full_section[len(red_flag_heading_keyword):].strip()

                                    if match_next_std_section:
                                        red_flag_content_only = red_flags_full_section[len(red_flag_heading_keyword) : len(red_flag_heading_keyword) + match_next_std_section.start()].strip()
                                        text_after_flags = red_flags_full_section[len(red_flag_heading_keyword) + match_next_std_section.start():].strip() # Corrected variable
                                    else:
                                        text_after_flags = ""


                                    st.markdown(text_before_flags) # Display text before the red flags
                                    st.markdown(f"**{actual_red_flag_heading_text.upper().strip()}**") # Display "RED FLAGS / KEY RISKS:"

                                    if "no specific red flags identified" in red_flag_content_only.lower():
                                        st.info(red_flag_content_only)
                                    else:
                                        st.warning(red_flag_content_only) # Use st.warning for actual flags
                                    
                                    if text_after_flags: # Display text after red flags if any was separated
                                        st.markdown(text_after_flags)
                                else:
                                    st.markdown(summary_content) # Display full summary if no specific section found

                                st.markdown("--- \n*Full summary text for copying:*")
                                st.code(summary_content, language=None) # Provides a copy button
                            else:
                                st.caption("No summary available, summary was empty, or an error occurred in summarization.")
                else: # Fallback if essential columns are missing
                    st.warning("The results digest is missing expected columns ('parent_company_no', 'summary_of_findings') for detailed display. Showing raw table instead:")
                    st.dataframe(st.session_state.ch_last_df, use_container_width=True, height=250)


                st.markdown("---")
                if st.button("🧠 Generate AI Narrative from Digest", key="ch_narrative_button_ui_updated"):
                    if st.session_state.ch_last_df is not None and not st.session_state.ch_last_df.empty:
                        with st.spinner("Generating AI narrative..."):
                            df_for_narrative = st.session_state.ch_last_df.copy()
                            if 'summary_of_findings' in df_for_narrative.columns:
                                # Filter for valid summaries (not errors, not too short)
                                valid_summaries_df = df_for_narrative[
                                    ~df_for_narrative['summary_of_findings'].astype(str).str.contains("No processable|Error:|No significant text|not configured|No content provided", case=False, na=False) &
                                    (df_for_narrative['summary_of_findings'].astype(str).str.len() > 50)
                                ]
                            else: valid_summaries_df = pd.DataFrame() # Empty if column missing

                            if not valid_summaries_df.empty:
                                relevant_cols_for_narrative = ["parent_company_no", "summary_of_findings"]
                                # Ensure columns exist before trying to use them
                                existing_cols_for_narrative = [col for col in relevant_cols_for_narrative if col in valid_summaries_df.columns]

                                if existing_cols_for_narrative:
                                    data_for_prompt_str = valid_summaries_df[existing_cols_for_narrative].to_json(orient="records", indent=2)
                                    # Limit length of data string sent to AI for narrative
                                    max_chars_for_narrative_data = 30000
                                    data_for_prompt_str = data_for_prompt_str[:max_chars_for_narrative_data] if len(data_for_prompt_str) > max_chars_for_narrative_data else data_for_prompt_str
                                    
                                    # System prompt for narrative generation
                                    sys_prompt_narrative = (
                                        "You are an expert financial analyst. Based on the following JSON data, where each entry contains a company number "
                                        "and its AI-generated summary of findings from Companies House documents, provide a concise, high-level narrative. "
                                        "Synthesize the key themes, risks (especially any explicitly mentioned as 'Red Flags'), and financial trends. "
                                        "If multiple companies are present, offer comparative points where appropriate. If only one company, focus on its standalone analysis. "
                                        "Avoid simply restating the input summaries; instead, draw conclusions and provide an overarching perspective. "
                                        "Highlight any common 'Red Flags' if they appear across multiple summaries or are significant for a single company."
                                    )
                                    user_prompt_narrative = f"DATA (JSON format):\n{data_for_prompt_str}\n\nNARRATIVE:"
                                    
                                    narr_text_resp = "Error: Narrative generation failed."
                                    try:
                                        narr_model = st.session_state.consult_digest_model # Use the general consultation model
                                        if not narr_model: raise ValueError("No AI model selected for narrative generation.")

                                        if narr_model.startswith("gpt-"):
                                            client = config.get_openai_client(); assert client
                                            resp = client.chat.completions.create(
                                                model=narr_model, temperature=0.25, max_tokens=1500, # Increased max_tokens for narrative
                                                messages=[{"role":"system","content":sys_prompt_narrative},{"role":"user","content":user_prompt_narrative}]
                                            )
                                            narr_text_resp = resp.choices[0].message.content.strip()
                                        elif narr_model.startswith("gemini-"):
                                            client = config.get_gemini_model(narr_model); assert client and config.genai
                                            full_prompt_narrative = f"{sys_prompt_narrative}\n\n{user_prompt_narrative}"
                                            resp = client.generate_content(
                                                full_prompt_narrative,
                                                generation_config=config.genai.types.GenerationConfig(temperature=0.25, max_output_tokens=1500) # Increased
                                            )
                                            # Handle potential blocking for Gemini narrative
                                            if hasattr(resp, 'text') and resp.text:
                                                narr_text_resp = resp.text.strip()
                                            elif hasattr(resp, 'prompt_feedback') and resp.prompt_feedback.block_reason:
                                                block_reason_narr = config.genai.types.BlockedReason(resp.prompt_feedback.block_reason).name
                                                narr_text_resp = f"Error: Narrative generation by Gemini was blocked (Reason: {block_reason_narr})."
                                                logger.error(f"Gemini narrative blocked: {block_reason_narr}. Feedback: {resp.prompt_feedback}")
                                            else:
                                                narr_text_resp = "Error: Gemini narrative response was empty or malformed."
                                                logger.error(f"Gemini narrative empty/malformed: {resp}")
                                        st.session_state.ch_last_narrative = narr_text_resp
                                    except Exception as e_narr:
                                        st.warning(f"Narrative generation failed: {e_narr}")
                                        logger.error(f"Narrative generation exception: {e_narr}", exc_info=True)
                                        st.session_state.ch_last_narrative = "Error during narrative generation."
                                else: # No relevant columns found in the valid_summaries_df
                                    st.info("No relevant columns (parent_company_no, summary_of_findings) in the filtered data for narrative generation.")
                                    st.session_state.ch_last_narrative = None
                            else: # No valid summaries met the criteria
                                st.info("No valid summaries found in the CH analysis results to generate a narrative.")
                                st.session_state.ch_last_narrative = None
                        st.rerun() # Rerun to display the narrative or message
                    else: # No CH DataFrame loaded
                        st.info("No Companies House data available to generate a narrative.")
                        st.session_state.ch_last_narrative = None


                if st.session_state.ch_last_narrative:
                    st.markdown("##### AI Narrative Summary of Digest"); st.markdown(st.session_state.ch_last_narrative)

            # This condition handles the case where run_ch_analysis_button_ui was true, but st.session_state.ch_last_df is None or empty
            elif run_ch_analysis_button_ui:
                st.warning("CH Analysis run initiated, but no data was generated or an error occurred before results could be displayed.")
                if st.session_state.ch_last_batch_metrics: # Check if metrics exist even if df is empty
                    if st.session_state.ch_last_batch_metrics.get("error"):
                        st.error(f"Pipeline Error: {st.session_state.ch_last_batch_metrics['error']}")
                    if st.session_state.ch_last_batch_metrics.get("notes"):
                         st.info(f"Pipeline Notes: {st.session_state.ch_last_batch_metrics['notes']}")


with tab_about_rendered:
    render_about_page()

# --- End of Main App Area ---